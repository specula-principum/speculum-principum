"""Tests for the DOCX parser implementation."""

from __future__ import annotations

from pathlib import Path

import pytest

from docx import Document as DocxBuilder  # type: ignore[import]

from src.parsing import registry
from src.parsing.base import ParseTarget, ParserError
from src.parsing.docx import DocxParser, docx_parser


def _build_sample_docx(path: Path) -> None:
    document = DocxBuilder()
    document.add_heading("Sample Document", level=1)
    document.add_paragraph("First paragraph of content.")
    document.add_paragraph("Second paragraph with more text.")

    document.add_paragraph("Item one", style="List Bullet")
    document.add_paragraph("Item two", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Row 1"
    table.cell(1, 1).text = "Row 1B"

    document.core_properties.author = "Test Author"
    document.core_properties.title = "Docx Parser Fixture"

    document.save(str(path))


def test_docx_parser_extracts_content_and_metadata(tmp_path) -> None:
    docx_path = tmp_path / "sample.docx"
    _build_sample_docx(docx_path)

    parser = DocxParser()
    target = ParseTarget(source=str(docx_path))

    assert parser.detect(target)

    document = parser.extract(target)

    assert "Sample Document" in document.segments[0]
    assert any(segment.startswith("- ") for segment in document.segments)
    assert any("| Header A | Header B |" in segment for segment in document.segments)

    metadata = document.metadata
    assert metadata["paragraph_count"] >= 2
    assert metadata["table_count"] == 1
    assert metadata["author"] == "Test Author"
    assert metadata["title"] == "Docx Parser Fixture"

    markdown = parser.to_markdown(document)
    assert "Sample Document" in markdown
    assert "| Header A | Header B |" in markdown


def test_docx_parser_registration(tmp_path) -> None:
    docx_path = tmp_path / "registry.docx"
    _build_sample_docx(docx_path)

    target = ParseTarget(source=str(docx_path))
    parser = registry.require_parser(target)

    assert parser.name == docx_parser.name
    document = parser.extract(target)
    assert document.segments


def test_docx_parser_handles_missing_file(tmp_path) -> None:
    parser = DocxParser()
    missing = tmp_path / "missing.docx"
    target = ParseTarget(source=str(missing))

    with pytest.raises(ParserError):
        parser.extract(target)


def test_docx_parser_warns_on_empty_doc(tmp_path) -> None:
    docx_path = tmp_path / "empty.docx"
    document = DocxBuilder()
    document.save(str(docx_path))

    parser = DocxParser()
    target = ParseTarget(source=str(docx_path))

    parsed = parser.extract(target)

    assert parsed.is_empty()
    assert any("no extractable content" in warning.lower() for warning in parsed.warnings)
